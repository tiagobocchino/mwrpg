"""
Gerador de .pdf e .docx a partir do .md fonte — ferramenta de suporte à
documentação de projeto, não faz parte do runtime do PushProcessos (por
isso vive em docs/entregaveis/_ferramentas/, com seu próprio
requirements.txt, separado do app).

Cada .md fonte tem um front-matter simples (chave: valor, entre linhas
"---") com metadado do documento (título, versão, autor, data). O corpo
é Markdown comum, com numeração de seção já escrita no próprio texto dos
títulos (ex: "## 2.1 Nome da seção") — assim a numeração fica idêntica
nos 3 formatos sem lógica de renderização separada.

Uso:
    python gerar_docs.py caminho/do/documento.md

Gera documento.pdf e documento.docx ao lado do .md, com capa, sumário,
numeração (a que já está no texto), rodapé com versão/data/página, e
imagens (inclusive os diagramas Mermaid pré-renderizados em
docs/entregaveis/diagramas/*.png) embutidas de verdade — nunca bloco de
código cru no lugar de diagrama.
"""

from __future__ import annotations

import re
import sys
from pathlib import Path

# Console do Windows usa cp1252 por padrão — os "✓" do relatório de progresso
# quebram sem isso (UnicodeEncodeError), então força utf-8 na saída.
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8")

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from markdown import markdown as md_to_html
from PIL import Image as PILImage
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfgen.canvas import Canvas
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    HRFlowable,
    Image,
    ListFlowable,
    ListItem,
    NextPageTemplate,
    PageBreak,
    PageTemplate,
    Paragraph,
    Preformatted,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.platypus.tableofcontents import TableOfContents

AZUL_ESCURO = "#0b1120"
AZUL_ACENTO = "#1d4ed8"
CINZA = "#555555"

# ─── Front-matter ───────────────────────────────────────────────────────────

def ler_documento(caminho: Path) -> tuple[dict, str]:
    texto = caminho.read_text(encoding="utf-8")
    if not texto.startswith("---"):
        raise ValueError("Documento sem front-matter (precisa começar com '---')")
    _, bloco_meta, corpo = texto.split("---", 2)
    meta: dict[str, str] = {}
    for linha in bloco_meta.strip().splitlines():
        if ":" not in linha:
            continue
        chave, valor = linha.split(":", 1)
        meta[chave.strip()] = valor.strip().strip('"')
    return meta, corpo.strip()


def _html_do_corpo(corpo_md: str) -> BeautifulSoup:
    html = md_to_html(corpo_md, extensions=["tables", "fenced_code", "sane_lists"])
    return BeautifulSoup(html, "html.parser")


# ─── DOCX ───────────────────────────────────────────────────────────────────

def _add_toc_field(paragraph) -> None:
    """Insere o campo de sumário nativo do Word (TOC \\o "1-3" \\h \\z \\u)."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-4" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    texto_aviso = OxmlElement("w:t")
    texto_aviso.text = "Clique com o botão direito e escolha \"Atualizar campo\" para gerar o sumário."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r_element = run._r
    r_element.append(fld_begin)
    r_element.append(instr)
    r_element.append(fld_sep)
    r_element.append(texto_aviso)
    r_element.append(fld_end)


def _docx_add_inline(paragraph, node) -> None:
    """Percorre filhos inline (texto, <strong>, <em>, <code>, <a>) de um nó."""
    for child in node.children:
        if isinstance(child, str):
            paragraph.add_run(child)
        elif child.name == "strong":
            r = paragraph.add_run(child.get_text())
            r.bold = True
        elif child.name == "em":
            r = paragraph.add_run(child.get_text())
            r.italic = True
        elif child.name == "code":
            r = paragraph.add_run(child.get_text())
            r.font.name = "Consolas"
            r.font.size = Pt(9.5)
        elif child.name == "a":
            r = paragraph.add_run(f"{child.get_text()} ({child.get('href', '')})")
            r.italic = True
        else:
            _docx_add_inline(paragraph, child)


def _docx_add_table(doc, node) -> None:
    linhas = node.find_all("tr")
    if not linhas:
        return
    n_cols = len(linhas[0].find_all(["td", "th"]))
    tabela = doc.add_table(rows=0, cols=n_cols)
    tabela.style = "Light Grid Accent 1"
    for linha in linhas:
        celulas_html = linha.find_all(["td", "th"])
        linha_docx = tabela.add_row()
        for i, celula in enumerate(celulas_html):
            linha_docx.cells[i].text = celula.get_text().strip()


def _docx_render_corpo(doc, soup: BeautifulSoup, pasta_base: Path) -> None:
    for node in soup.find_all(recursive=False):
        if node.name in ("h1", "h2", "h3", "h4"):
            nivel = int(node.name[1])
            doc.add_heading(node.get_text(), level=nivel)
        elif node.name == "p":
            if node.find("img"):
                img = node.find("img")
                caminho_img = pasta_base / img["src"]
                if caminho_img.exists():
                    largura_px, _ = PILImage.open(caminho_img).size
                    largura = Inches(6) if largura_px > 600 else Inches(4.5)
                    doc.add_picture(str(caminho_img), width=largura)
                    legenda = img.get("alt", "")
                    if legenda:
                        p = doc.add_paragraph(legenda)
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        p.runs[0].italic = True
                        p.runs[0].font.size = Pt(9)
            else:
                p = doc.add_paragraph()
                _docx_add_inline(p, node)
        elif node.name in ("ul", "ol"):
            estilo = "List Bullet" if node.name == "ul" else "List Number"
            for li in node.find_all("li", recursive=False):
                p = doc.add_paragraph(style=estilo)
                _docx_add_inline(p, li)
        elif node.name == "pre":
            codigo = node.get_text()
            p = doc.add_paragraph()
            r = p.add_run(codigo)
            r.font.name = "Consolas"
            r.font.size = Pt(8.5)
            p.paragraph_format.left_indent = Cm(0.5)
        elif node.name == "blockquote":
            p = doc.add_paragraph()
            _docx_add_inline(p, node)
            for r in p.runs:
                r.italic = True
        elif node.name == "table":
            _docx_add_table(doc, node)
        elif node.name == "hr":
            doc.add_paragraph("─" * 40).alignment = WD_ALIGN_PARAGRAPH.CENTER


def gerar_docx(meta: dict, corpo_md: str, caminho_saida: Path, pasta_base: Path) -> None:
    doc = Document()

    estilo_normal = doc.styles["Normal"]
    estilo_normal.font.name = "Calibri"
    estilo_normal.font.size = Pt(11)

    # ── Capa ──
    titulo = doc.add_paragraph()
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = titulo.add_run(meta.get("titulo", ""))
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)

    if meta.get("subtitulo"):
        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = sub.add_run(meta["subtitulo"])
        r.font.size = Pt(14)
        r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    for _ in range(6):
        doc.add_paragraph()

    tabela_meta = doc.add_table(rows=0, cols=2)
    for chave, rotulo in [
        ("projeto", "Projeto"), ("codigo", "Documento nº"), ("versao", "Versão"),
        ("autor", "Autoria"), ("data", "Data"),
    ]:
        if meta.get(chave):
            linha = tabela_meta.add_row()
            linha.cells[0].text = rotulo
            linha.cells[1].text = meta[chave]
            linha.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_page_break()

    # ── Sumário ──
    doc.add_heading("Sumário", level=1)
    _add_toc_field(doc.add_paragraph())
    doc.add_page_break()

    # ── Corpo ──
    soup = _html_do_corpo(corpo_md)
    _docx_render_corpo(doc, soup, pasta_base)

    # ── Rodapé ──
    secao = doc.sections[0]
    rodape = secao.footer.paragraphs[0]
    rodape.text = f"{meta.get('projeto', '')} — {meta.get('titulo', '')} — v{meta.get('versao', '')}"
    rodape.style = doc.styles["Normal"]
    for r in rodape.runs:
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.save(caminho_saida)


# ─── PDF ────────────────────────────────────────────────────────────────────

def _estilos_pdf():
    ss = getSampleStyleSheet()
    ss.add(ParagraphStyle("CapaTitulo", parent=ss["Title"], fontSize=26, textColor=colors.HexColor(AZUL_ACENTO), spaceAfter=14))
    ss.add(ParagraphStyle("CapaSubtitulo", parent=ss["Normal"], fontSize=14, textColor=colors.HexColor(CINZA), alignment=1, spaceAfter=6))
    ss.add(ParagraphStyle("CapaMeta", parent=ss["Normal"], fontSize=11, alignment=1, spaceAfter=4))
    ss.add(ParagraphStyle("H1", parent=ss["Heading1"], fontSize=17, spaceBefore=18, spaceAfter=8, textColor=colors.HexColor(AZUL_ESCURO)))
    ss.add(ParagraphStyle("H2", parent=ss["Heading2"], fontSize=14, spaceBefore=14, spaceAfter=6, textColor=colors.HexColor(AZUL_ACENTO)))
    ss.add(ParagraphStyle("H3", parent=ss["Heading3"], fontSize=12, spaceBefore=10, spaceAfter=4))
    ss.add(ParagraphStyle("H4", parent=ss["Heading4"], fontSize=11, spaceBefore=8, spaceAfter=4))
    ss.add(ParagraphStyle("CorpoTexto", parent=ss["BodyText"], fontSize=10.5, leading=15, spaceAfter=8))
    ss.add(ParagraphStyle("ListaTexto", parent=ss["BodyText"], fontSize=10.5, leading=15, spaceAfter=4, leftIndent=16))
    ss.add(ParagraphStyle("Citacao", parent=ss["BodyText"], fontSize=10.5, leftIndent=20, textColor=colors.HexColor(CINZA), spaceAfter=8))
    ss.add(ParagraphStyle("Legenda", parent=ss["Normal"], fontSize=9, alignment=1, textColor=colors.HexColor(CINZA), spaceAfter=10))
    ss.add(ParagraphStyle("SumarioTitulo", parent=ss["Heading1"], fontSize=17))
    return ss


def _escapar(texto: str) -> str:
    return texto.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _pdf_inline_para_texto(node) -> str:
    partes = []
    for child in node.children:
        if isinstance(child, str):
            partes.append(_escapar(child))
        elif child.name == "strong":
            partes.append(f"<b>{_escapar(child.get_text())}</b>")
        elif child.name == "em":
            partes.append(f"<i>{_escapar(child.get_text())}</i>")
        elif child.name == "code":
            partes.append(f"<font face='Courier'>{_escapar(child.get_text())}</font>")
        elif child.name == "a":
            partes.append(f"<i>{_escapar(child.get_text())} ({_escapar(child.get('href', ''))})</i>")
        else:
            partes.append(_pdf_inline_para_texto(child))
    return "".join(partes)


class _DocComSumario(BaseDocTemplate):
    """Doc template com 2 passes — permite número de página real no sumário."""

    def afterFlowable(self, flowable):
        if isinstance(flowable, Paragraph):
            estilo = flowable.style.name
            texto = flowable.getPlainText()
            if estilo == "H1":
                self.notify("TOCEntry", (0, texto, self.page))
                self.canv.bookmarkPage(texto)
                self.canv.addOutlineEntry(texto, texto, level=0)
            elif estilo == "H2":
                self.notify("TOCEntry", (1, texto, self.page))
                self.canv.bookmarkPage(texto)
                self.canv.addOutlineEntry(texto, texto, level=1)


def _pdf_render_corpo(soup: BeautifulSoup, pasta_base: Path, ss, largura_util: float) -> list:
    elementos = []
    for node in soup.find_all(recursive=False):
        if node.name in ("h1", "h2", "h3", "h4"):
            elementos.append(Paragraph(_escapar(node.get_text()), ss[node.name.upper()]))
        elif node.name == "p":
            if node.find("img"):
                img = node.find("img")
                caminho_img = pasta_base / img["src"]
                if caminho_img.exists():
                    largura_px, altura_px = PILImage.open(caminho_img).size
                    razao = altura_px / largura_px
                    largura = min(largura_util, 14 * cm)
                    altura = largura * razao
                    # imagem alta (diagrama vertical) estoura a página se só a
                    # largura for limitada — limita a altura também e recalcula
                    # a largura pela mesma razão, senão o reportlab quebra com
                    # LayoutError em vez de só encolher a imagem.
                    altura_max = 23 * cm
                    if altura > altura_max:
                        altura = altura_max
                        largura = altura / razao
                    elementos.append(Spacer(1, 8))
                    elementos.append(Image(str(caminho_img), width=largura, height=altura, hAlign="CENTER"))
                    legenda = img.get("alt", "")
                    if legenda:
                        elementos.append(Paragraph(_escapar(legenda), ss["Legenda"]))
            else:
                elementos.append(Paragraph(_pdf_inline_para_texto(node), ss["CorpoTexto"]))
        elif node.name == "ul":
            itens = [
                ListItem(Paragraph(_pdf_inline_para_texto(li), ss["CorpoTexto"]), leftIndent=14)
                for li in node.find_all("li", recursive=False)
            ]
            elementos.append(ListFlowable(itens, bulletType="bullet"))
        elif node.name == "ol":
            # numeração manual em vez do bulletType="1" do ListFlowable: em
            # certas combinações de fonte/reportlab ele quebra tentando tratar
            # o número do marcador como bytes (AttributeError: 'int' object
            # has no attribute 'decode') — numerar no próprio texto é robusto
            # e mantém a mesma numeração explícita já usada nos títulos.
            for i, li in enumerate(node.find_all("li", recursive=False), start=1):
                elementos.append(Paragraph(f"{i}. {_pdf_inline_para_texto(li)}", ss["ListaTexto"]))
        elif node.name == "pre":
            elementos.append(Preformatted(node.get_text(), ss["Code"] if "Code" in ss else ss["Normal"]))
        elif node.name == "blockquote":
            elementos.append(Paragraph(_pdf_inline_para_texto(node), ss["Citacao"]))
        elif node.name == "table":
            linhas = node.find_all("tr")
            dados = [[c.get_text().strip() for c in linha.find_all(["td", "th"])] for linha in linhas]
            tabela = Table(dados, hAlign="LEFT")
            tabela.setStyle(TableStyle([
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#e8edfa")),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]))
            elementos.append(Spacer(1, 6))
            elementos.append(tabela)
            elementos.append(Spacer(1, 6))
        elif node.name == "hr":
            elementos.append(HRFlowable(width="100%", color=colors.HexColor("#cccccc")))
    return elementos


def _rodape_pdf(meta: dict):
    def _desenhar(canvas: Canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 7.5)
        canvas.setFillColor(colors.HexColor("#999999"))
        texto = f"{meta.get('projeto', '')} — {meta.get('titulo', '')} — v{meta.get('versao', '')}"
        canvas.drawString(2.5 * cm, 1.2 * cm, texto)
        canvas.drawRightString(A4[0] - 2 * cm, 1.2 * cm, f"página {doc.page}")
        canvas.restoreState()
    return _desenhar


def gerar_pdf(meta: dict, corpo_md: str, caminho_saida: Path, pasta_base: Path) -> None:
    ss = _estilos_pdf()
    if "Code" not in ss:
        ss.add(ParagraphStyle("Code", parent=ss["Normal"], fontName="Courier", fontSize=8.5, leftIndent=10, backColor=colors.HexColor("#f2f2f2")))

    largura_pagina = A4[0] - 2.5 * cm - 2 * cm
    doc = _DocComSumario(
        str(caminho_saida), pagesize=A4,
        topMargin=2 * cm, bottomMargin=2.2 * cm, leftMargin=2.5 * cm, rightMargin=2 * cm,
    )
    frame = Frame(doc.leftMargin, doc.bottomMargin, doc.width, doc.height, id="corpo")
    doc.addPageTemplates([PageTemplate(id="corpo", frames=[frame], onPage=_rodape_pdf(meta))])

    elementos = []

    # ── Capa ──
    elementos.append(Spacer(1, 6 * cm))
    elementos.append(Paragraph(_escapar(meta.get("titulo", "")), ss["CapaTitulo"]))
    if meta.get("subtitulo"):
        elementos.append(Paragraph(_escapar(meta["subtitulo"]), ss["CapaSubtitulo"]))
    elementos.append(Spacer(1, 3 * cm))
    for chave, rotulo in [
        ("projeto", "Projeto"), ("codigo", "Documento nº"), ("versao", "Versão"),
        ("autor", "Autoria"), ("data", "Data"),
    ]:
        if meta.get(chave):
            elementos.append(Paragraph(f"<b>{rotulo}:</b> {_escapar(meta[chave])}", ss["CapaMeta"]))
    elementos.append(PageBreak())

    # ── Sumário ──
    elementos.append(Paragraph("Sumário", ss["SumarioTitulo"]))
    sumario = TableOfContents()
    sumario.levelStyles = [
        ParagraphStyle(name="TOC1", fontSize=11, leading=16, firstLineIndent=0),
        ParagraphStyle(name="TOC2", fontSize=10, leading=14, leftIndent=14),
    ]
    elementos.append(sumario)
    elementos.append(PageBreak())

    # ── Corpo ──
    soup = _html_do_corpo(corpo_md)
    elementos.extend(_pdf_render_corpo(soup, pasta_base, ss, largura_pagina))

    doc.multiBuild(elementos)


# ─── CLI ────────────────────────────────────────────────────────────────────

def gerar(caminho_md: str) -> None:
    caminho = Path(caminho_md).resolve()
    meta, corpo_md = ler_documento(caminho)
    pasta_base = caminho.parent

    caminho_docx = caminho.with_suffix(".docx")
    caminho_pdf = caminho.with_suffix(".pdf")

    gerar_docx(meta, corpo_md, caminho_docx, pasta_base)
    print(f"✓ {caminho_docx.name}")

    gerar_pdf(meta, corpo_md, caminho_pdf, pasta_base)
    print(f"✓ {caminho_pdf.name}")


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Uso: python gerar_docs.py caminho/do/documento.md [outro.md ...]")
        sys.exit(1)
    for arg in sys.argv[1:]:
        print(f"\n=== {arg} ===")
        gerar(arg)
